from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Iterable


@dataclass(frozen=True)
class ReportSection:
    title: str
    content: str


def build_report_markdown(
    report_title: str,
    document_name: str,
    sections: Iterable[ReportSection],
) -> str:
    section_blocks = []
    for section in sections:
        section_blocks.append(f"## {section.title}\n\n{section.content.strip()}")

    body = "\n\n".join(section_blocks).strip()
    return (
        f"# {report_title.strip() or 'InsightForge Report'}\n\n"
        f"**Source document:** {document_name}\n\n"
        f"{body}"
    ).strip()


def build_report_docx(
    report_title: str,
    document_name: str,
    sections: Iterable[ReportSection],
) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX export requires python-docx. Install it with 'pip install python-docx'."
        ) from exc

    document = Document()
    document.add_heading(report_title.strip() or "InsightForge Report", level=0)
    document.add_paragraph(f"Source document: {document_name}")

    for section in sections:
        document.add_heading(section.title, level=1)
        for line in _content_lines(section.content):
            if line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_report_pdf(
    report_title: str,
    document_name: str,
    sections: Iterable[ReportSection],
) -> bytes:
    try:
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError(
            "PDF export requires reportlab. Install it with 'pip install reportlab'."
        ) from exc

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=(8.5 * inch, 11 * inch),
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    heading_style = styles["Heading2"]
    body_style = ParagraphStyle(
        "ReportBody",
        parent=styles["BodyText"],
        leading=16,
        spaceAfter=8,
    )
    bullet_style = ParagraphStyle(
        "ReportBullet",
        parent=body_style,
        leftIndent=14,
        bulletIndent=0,
    )

    story = [
        Paragraph(report_title.strip() or "InsightForge Report", title_style),
        Spacer(1, 0.18 * inch),
        Paragraph(f"Source document: {document_name}", body_style),
        Spacer(1, 0.15 * inch),
    ]

    for section in sections:
        story.append(Paragraph(section.title, heading_style))
        story.append(Spacer(1, 0.08 * inch))
        for line in _content_lines(section.content):
            if line.startswith("- "):
                story.append(
                    Paragraph(_escape_pdf_text(line[2:]), bullet_style, bulletText="•")
                )
            else:
                story.append(Paragraph(_escape_pdf_text(line), body_style))
        story.append(Spacer(1, 0.12 * inch))

    document.build(story)
    return buffer.getvalue()


def _content_lines(content: str) -> list[str]:
    lines = []
    for raw_line in content.strip().splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("#"):
            line = line.lstrip("#").strip()
        lines.append(line)
    return lines or ["No content provided."]


def _escape_pdf_text(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
